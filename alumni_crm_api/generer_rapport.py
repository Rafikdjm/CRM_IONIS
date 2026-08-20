#!/usr/bin/env python3
"""
Script pour générer le document Word de préparation à l'oral de stage.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime


def set_cell_shading(cell, color):
    """Applique un fond coloré à une cellule."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color,
        qn("w:val"): "clear",
    })
    shading.append(shading_elm)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Ajoute un tableau stylisé au document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # En-têtes
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "2B579A")

    # Données
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if row_idx % 2 == 0:
                set_cell_shading(cell, "E8EFF7")

    return table


def generer_document():
    doc = Document()

    # ── Style de base ──
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # ═══════════════════════════════════════════════════════════
    # PAGE DE GARDE
    # ═══════════════════════════════════════════════════════════
    doc.add_paragraph("\n\n\n")
    title = doc.add_heading("Rapport de Préparation à l'Oral de Stage", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("CRM Alumni — Conception d'un outil centralisé de gestion des données anciens élèves")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(43, 87, 154)

    doc.add_paragraph("\n\n")
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"Document généré le {datetime.date.today().strftime('%d/%m/%Y')}")
    run.font.size = Pt(10)
    run.font.italic = True

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # TABLE DES MATIÈRES
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("Table des matières", level=1)
    toc_items = [
        "1. Contexte précis du stage et objectifs",
        "2. Missions confiées et responsabilités attribuées",
        "3. Résultats obtenus, impact des actions menées",
        "4. Réponse à la problématique initiale",
        "5. Organisation du travail en équipe, ressources",
        "6. Méthodes et stratégies mises en œuvre",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 1. CONTEXTE PRÉCIS DU STAGE ET OBJECTIFS
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("1. Contexte précis du stage et objectifs", level=1)

    doc.add_heading("1.1 Durée et période du stage", level=2)
    doc.add_paragraph(
        "Le projet a été développé dans le cadre d'un stage de formation, "
        "réalisé au sein de l'établissement scolaire. La durée du stage n'est pas "
        "explicitement renseignée dans le code source, mais l'évolution progressive "
        "du projet — visible à travers les 5 migrations SQL numérotées (001 à 006, "
        "avec un numéro 003 absent), la refactorisation de l'architecture en modules, "
        "et la structuration du README en « 4 axes de la revue » — témoigne d'un "
        "travail mené sur plusieurs semaines avec des livrables itératifs."
    )

    doc.add_heading("1.2 Présentation du sujet", level=2)
    doc.add_paragraph(
        "Le sujet de stage a été présenté dans le contexte de la gestion des anciens "
        "élèves (alumni) de l'établissement. Le README du projet mentionne un "
        "« rapport de stage », confirmant la dimension formelle de la mission. "
        "Le sujet initial a été formulé sous la forme d'une problématique : "
        "comment concevoir un outil centralisé de gestion des données alumni "
        "tout en garantissant la conformité RGPD et un suivi du cycle de vie des données."
    )

    doc.add_heading("1.3 Objectif business", level=2)
    doc.add_paragraph("L'objectif business, au-delà de la fiche sujet, était de :")
    objectives = [
        "Centraliser les données des anciens élèves dans une base de données relationnelle unique, remplaçant les fichiers Excel dispersés",
        "Suivre le parcours professionnel des diplômés : expériences professionnelles, salaires, certifications, secteurs d'activité",
        "Automatiser le calcul des indicateurs d'insertion professionnelle (taux d'emploi, salaire moyen, taux d'insertion à 6 mois)",
        "Administrer des questionnaires annuels pour collecter des données actualisées auprès des alumni",
        "Garantir la conformité RGPD avec un système de consentement, d'anonymisation et de journalisation (audit trail)",
        "Fournir des outils d'import/export pour faciliter la mise à jour et l'exploitation des données",
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style="List Bullet")

    # ═══════════════════════════════════════════════════════════
    # 2. MISSIONS CONFIÉES ET RESPONSABILITÉS
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("2. Missions confiées et responsabilités attribuées", level=1)

    doc.add_heading("2.1 Double volet : Management + Informatique", level=2)
    doc.add_paragraph(
        "Le projet couvre les deux volets, avec une dominante informatique. "
        "Le volet management se manifeste dans la gouvernance des données, "
        "la conception des processus de mise à jour, et la réflexion sur la "
        "conformité RGPD. Le volet informatique est le cœur du projet avec "
        "le développement complet de l'API."
    )

    doc.add_heading("2.2 Réalisations concrètes — ce qui me revient en propre", level=2)

    doc.add_heading("Modélisation des données", level=3)
    doc.add_paragraph(
        "Conception du modèle relationnel complet de la base de données, "
        "depuis le MCD jusqu'au SQL PostgreSQL. Le schéma comprend 11 tables "
        "(voir section 3 pour le détail). Les migrations sont numérotées et "
        "traquées via une table schema_migrations, démontrant une démarche "
        "structurée d'évolution de la base."
    )

    doc.add_heading("Conformité RGPD", level=3)
    doc.add_paragraph(
        "Conception et implémentation du système de gestion des consentements RGPD :"
    )
    rgpd_features = [
        "Table CONSENTEMENT_RGPD avec champs : date, type (ex: PROFIL_ALUMNI), statut (ACCORD/REFUS), canal de collecte",
        "Fonctionnalité d'anonymisation logique : les données personnelles sont masquées (nom → « ANONYMISE », email → « ANONYMISE_{id}@anonymise.io ») sans suppression physique, préservant l'intégrité référentielle",
        "Anonymisation des expériences professionnelles liées (intitulé → « ANONYMISE », salaire → 0)",
        "Journal d'audit (AUDIT_LOG) pour tracer toutes les opérations de nettoyage et d'archivage",
        "Vérification que l'étudiant n'a pas re-granté son consentement avant anonymisation",
    ]
    for feat in rgpd_features:
        doc.add_paragraph(feat, style="List Bullet")

    doc.add_heading("Développement de l'API REST", level=3)
    doc.add_paragraph(
        "Développement complet de l'API avec FastAPI, comprenant :"
    )
    api_features = [
        "30+ endpoints RESTful couvrant toutes les entités métier",
        "Authentification double : API Key pour l'admin + OTP par email pour les alumni (JWT)",
        "Rate limiting en mémoire sur les endpoints d'authentification",
        "Pagination systématique sur tous les endpoints de listing",
        "Validation des données avec Pydantic v2",
        "Import/Export Excel (openpyxl) avec template personnalisé",
        "Import CSV/Excel via Pandas avec préchargement optimisé des entreprises",
        "Système de questionnaires annuels avec réponses stockées en JSONB",
        "Indicateurs d'insertion professionnelle (taux d'emploi, salaire moyen, taux à 6 mois)",
        "Nettoyage de la base : détection des orphelins et doublons avec mode dry-run",
    ]
    for feat in api_features:
        doc.add_paragraph(feat, style="List Bullet")

    doc.add_heading("Frontend React et corrections récentes", level=3)
    doc.add_paragraph(
        "En complément de l'API, un frontend React/Vite complet (dossier "
        "alumni_crm_front) a été développé et connecté à l'API, avec les "
        "corrections et ajouts récents suivants :"
    )
    recent_features = [
        "Correction du bug d'authentification RGPD : le token admin et le token alumni partageaient la même clé de stockage navigateur, ce qui pouvait envoyer un JWT admin sur les routes /rgpd/* (erreur 403). Correction : vérification du rôle dans le payload du JWT avant les appels sensibles (ensureAlumniToken), purge automatique du token en cas de session orpheline (intercepteur 401) et clés de session distinctes (admin_role vs alumni_id).",
        "Actions groupées (bulk) dans l'interface admin des demandes RGPD : traiter, rejeter (avec motif), exporter et supprimer en masse, avec modales de confirmation et affichage du nombre de demandes sélectionnées.",
        "Purge définitive différée RGPD : après anonymisation, les comptes peuvent être supprimés définitivement après un délai configurable (PURGE_DELAY_MONTHS, défaut 6 mois). Prévisualisation et exécution via les endpoints admin /admin/demandes-rgpd/purge-anonymises, script CLI purge.py, et journalisation complète dans AUDIT_LOG.",
        "Rate-limiting différencié selon l'environnement : le rate-limiting OTP est désactivé en développement (ENV=development) pour faciliter les tests, et appliqué strictement en production (par email et par IP).",
    ]
    for feat in recent_features:
        doc.add_paragraph(feat, style="List Bullet")

    doc.add_heading("2.3 Livrables intermédiaires", level=2)
    doc.add_paragraph(
        "Les livrables intermédiaires sont visibles dans l'historique des migrations :"
    )
    deliverables = [
        "Migration 001 : Table AUDIT_LOG + indexes → Mise en place de la traçabilité",
        "Migration 002 : Champs profil alumni étendus (adresse, ville, pays, LinkedIn, compétences JSONB, statut disponibilité)",
        "Migration 004 : Tables QUESTIONNAIRE, QUESTION, REPONSE_QUESTIONNAIRE → Système de questionnaires annuels",
        "Migration 005 : Colonne tag sur QUESTION → Identification des KPI",
        "Migration 006 : Colonne conditionnee_statut_emploi → Questions conditionnelles",
        "README structuré en 4 axes de la revue = documentation de suivi des améliorations",
    ]
    for d in deliverables:
        doc.add_paragraph(d, style="List Bullet")

    # ═══════════════════════════════════════════════════════════
    # 3. RÉSULTATS OBTENUS
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("3. Résultats obtenus, impact des actions menées", level=1)

    doc.add_heading("3.1 État d'avancement", level=2)
    doc.add_paragraph(
        "Le projet est à un stade de prototype fonctionnel complet :"
    )
    status_items = [
        "Schéma de BDD terminé : 11 tables, toutes les relations et contraintes en place",
        "API fonctionnelle : 30+ endpoints testables via Swagger UI (/docs)",
        "Authentification opérationnelle : OTP par email (mode console en dev, mode Resend en production) + auth admin par code",
        "Import/Export Excel fonctionnel avec template personnalisé",
        "Système de questionnaires annuels complet (CRUD + réponses + KPI)",
        "Frontend React/Vite développé et fonctionnel (dossier alumni_crm_front) : authentification OTP (AuthPage), espace alumni (profil, carrière, consentement RGPD, questionnaires) et espace admin (dashboard, annuaire avec filtres, import/export, gestion des questionnaires, demandes RGPD avec actions groupées)",
    ]
    for item in status_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3.2 Chiffres clés", level=2)

    add_styled_table(doc,
        ["Indicateur", "Valeur", "Détail"],
        [
            ["Tables BDD", "11", "PROMOTION, ETUDIANT, ENTREPRISE, EXPERIENCE_PRO, CERTIFICATION, OBTIENT, CONSENTEMENT_RGPD, QUESTIONNAIRE, QUESTION, REPONSE_QUESTIONNAIRE, AUDIT_LOG"],
            ["Tables supplémentaires", "2", "otp_codes (auth), schema_migrations (tracking)"],
            ["Endpoints API", "30+", "Répartis dans 12 routers"],
            ["Champs profil étudiant", "15", "nom, prenom, email, email_academique, telephone, date_naissance, parcours_anterieur, date_inscription, id_promotion, address, city, country, linkedin, availability_status, skills (JSONB)"],
            ["Fichiers migrations SQL", "5", "001_audit_log, 002_alumni_profile_fields, 004_questionnaire_annuel, 005_question_tag, 006_question_statut_emploi"],
            ["Pool de connexions", "5", "Configurable via DB_POOL_SIZE"],
            ["Taille max upload", "5 Mo", "Configurable via MAX_UPLOAD_SIZE_MB"],
            ["TTL code OTP", "10 minutes", "5 tentatives max"],
            ["Rate limit OTP", "3/email/10min, 10/IP/1h", "Protection contre le brute force"],
            ["Rate limit admin", "5/IP/10min", "Protection contre le brute force"],
        ]
    )

    doc.add_heading("3.3 Indicateurs calculables par l'API", level=2)
    doc.add_paragraph(
        "L'API calcule automatiquement les indicateurs suivants :"
    )
    indicators = [
        "Taux d'emploi par promotion : (étudiants en poste / total étudiants) × 100",
        "Salaire moyen par promotion",
        "Taux d'emploi à 6 mois : diplômés ayant au moins une expérience avant décembre de l'année d'obtention",
        "Nombre d'alumni actifs (ayant au moins une expérience enregistrée)",
        "Taux de réponse : pourcentage d'alumni ayant complété leur profil",
        "Répartition par secteur d'activité",
        "KPI par tag de question (ex: adéquation formation)",
    ]
    for ind in indicators:
        doc.add_paragraph(ind, style="List Bullet")

    doc.add_heading("3.4 Impact concret pour l'établissement", level=2)
    impacts = [
        "Centralisation : passage de fichiers Excel dispersés à une base de données relationnelle unique et cohérente",
        "Automatisation : import massif via Excel (2 méthodes : /import/excel et /upload-etudiants/), export complet en un clic",
        "Fiabilité : détection et suppression automatique des doublons (sur clé nom+prénom+email) et des enregistrements orphelins",
        "Indicateurs : calcul automatique des taux d'insertion, permettant de répondre aux obligations de reporting",
        "Conformité RGPD : anonymisation automatique en cas de refus de consentement, traçabilité complète via audit log",
        "Suivi alumni : questionnaires annuels avec stockage flexible (JSONB), questions conditionnelles selon le statut d'emploi",
        "Sécurité : authentification OTP + JWT pour les alumni, API Key + JWT pour l'admin, rate limiting, messages d'erreur sanitizés",
    ]
    for impact in impacts:
        doc.add_paragraph(impact, style="List Bullet")

    # ═══════════════════════════════════════════════════════════
    # 4. RÉPONSE À LA PROBLÉMATIQUE INITIALE
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("4. Réponse à la problématique initiale", level=1)

    doc.add_heading("4.1 Rappel de la problématique", level=2)
    doc.add_paragraph(
        "La problématique formulée était : « Comment concevoir un outil centralisé "
        "de gestion des données alumni tout en garantissant la conformité RGPD "
        "et un suivi du cycle de vie des données ? »"
    )

    doc.add_heading("4.2 Réponse point par point", level=2)

    doc.add_paragraph(
        "Le frontend React/Vite démontre concrètement l'interface utilisateur pour "
        "les deux profils exigés par le sujet de stage : l'espace admin (dashboard, "
        "annuaire avec filtres, import/export, gestion des questionnaires, demandes "
        "RGPD avec actions groupées) et l'espace alumni (profil, carrière, "
        "consentement RGPD, questionnaires). Chaque profil dispose ainsi d'une "
        "interface dédiée, connectée à ses endpoints respectifs et protégée par un "
        "contrôle d'accès par rôle (JWT admin vs JWT alumni)."
    )

    # Tableau de réponse
    add_styled_table(doc,
        ["Axe de la problématique", "Réponse apportée", "Fichier(s) concerné(s)"],
        [
            [
                "Centralisation des données",
                "Base de données PostgreSQL relationnelle avec 11 tables, API REST FastAPI avec 30+ endpoints, import/export Excel",
                "database.py, main.py, routers/*.py"
            ],
            [
                "Gouvernance des données",
                "Table AUDIT_LOG pour tracer chaque opération. Workflow de nettoyage (orphelins + doublons) avec mode dry-run. Table schema_migrations pour suivre l'évolution de la BDD",
                "routers/cleanup.py, run_migrations.py"
            ],
            [
                "Conformité RGPD",
                "Table CONSENTEMENT_RGPD (date, type, statut ACCORD/REFUS, canal). Anonymisation logique (UPDATE, pas DELETE) des données personnelles et salariales. Vérification du consentement avant archivage",
                "routers/rgpd.py, routers/cleanup.py"
            ],
            [
                "Cycle de vie des données",
                "Création (import/CUD) → Mise à jour (PUT/PATCH) → Collecte (questionnaires) → Analyse (KPI/indicateurs) → Archivage/anonymisation RGPD, le tout tracé par AUDIT_LOG",
                "Tous les routers"
            ],
            [
                "Exploitation des données",
                "Endpoints admin avec indicateurs : taux d'emploi, salaire moyen, taux à 6 mois, répartition sectorielle, KPI par tag de question. Filtrage avancé par promotion/secteur/entreprise",
                "routers/admin.py"
            ],
            [
                "Sécurité des données",
                "API Key pour les routes admin. OTP + JWT pour les alumni. Rate limiting. Messages d'erreur sanitizés. Variables d'environnement pour les secrets",
                "security.py, routers/otp.py, routers/admin_auth.py, config.py"
            ],
            [
                "Interfaces utilisateur distinctes (admin / alumni)",
                "Frontend React/Vite : espace admin et espace alumni séparés, chacun connecté à ses endpoints dédiés, avec contrôle d'accès par rôle (JWT admin vs JWT alumni)",
                "alumni_crm_front (React/Vite)"
            ],
        ]
    )

    # ═══════════════════════════════════════════════════════════
    # 5. ORGANISATION DU TRAVAIL
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("5. Organisation du travail en équipe, ressources", level=1)

    doc.add_heading("5.1 Travail en équipe", level=2)
    doc.add_paragraph(
        "Le projet a été réalisé principalement en solo, comme en témoigne "
        "la structure cohérente du dépôt (un seul ensemble de fichiers, pas de "
        "branching visible, architecture unifiée). Le README mentionne des "
        "corrections organisées en « 4 axes de la revue », ce qui suggère des "
        "points d'étape réguliers avec un tuteur pédagogique ou un référent technique."
    )

    doc.add_heading("5.2 Outils utilisés", level=2)

    add_styled_table(doc,
        ["Catégorie", "Outil", "Usage"],
        [
            ["Langage", "Python 3.13", "Runtime principal de l'application"],
            ["Framework web", "FastAPI", "Framework API REST async avec auto-documentation Swagger"],
            ["Serveur ASGI", "Uvicorn", "Serveur de développement et production"],
            ["Base de données", "PostgreSQL", "SGBD relationnel"],
            ["Driver BDD", "pg8000", "Driver Python pur PostgreSQL (sans ORM)"],
            ["Validation", "Pydantic v2", "Validation et sérialisation des données (avec EmailStr)"],
            ["Auth", "PyJWT", "Création et vérification de tokens JWT"],
            ["Email", "Resend API", "Envoi d'emails transactionnels (codes OTP)"],
            ["Import/Export", "Pandas + openpyxl", "Lecture de CSV/Excel pour import, génération d'Excel pour export"],
            ["Variables d'env", "python-dotenv", "Chargement des variables d'environnement depuis .env"],
            ["Migrations", "Système custom (run_migrations.py)", "Exécution séquentielle de fichiers SQL avec tracking"],
            ["Frontend (implémenté)", "React / Vite", "Application complète connectée à l'API : authentification OTP, espace alumni (profil, carrière, consentement, questionnaires) et espace admin (dashboard, annuaire, import/export, questionnaires, demandes RGPD)"],
        ]
    )

    doc.add_heading("5.3 Points d'étape", level=2)
    doc.add_paragraph(
        "La structure du README (« version corrigée », « 4 axes de la revue », "
        "« limites connues à garder en tête pour le rapport de stage ») indique "
        "des points d'étape réguliers avec un tuteur. Les axes de la revue "
        "couvrent :"
    )
    review_axes = [
        "Robustesse et sécurité (identifiants en variables d'environnement, protection des routes admin, gestion des erreurs, élimination des race conditions TOCTOU)",
        "Performance (pool de connexions, pagination, préchargement pour les imports)",
        "Bonnes pratiques (Pydantic v2, sérialisation par cursor.description, contraintes de validation)",
        "Maintenabilité (découpage du fichier unique en 12 routers domain-specific)",
    ]
    for axis in review_axes:
        doc.add_paragraph(axis, style="List Bullet")

    # ═══════════════════════════════════════════════════════════
    # 6. MÉTHODES ET STRATÉGIES
    # ═══════════════════════════════════════════════════════════
    doc.add_heading("6. Méthodes et stratégies mises en œuvre", level=1)

    doc.add_heading("6.1 Modélisation de la base de données", level=2)
    doc.add_paragraph(
        "La modélisation suit une approche relationnelle classique (MCD → MLD → SQL). "
        "Le schéma est conçu sans ORM : les requêtes SQL sont écrites directement "
        "dans les routers, ce qui démontre une compréhension profonde du modèle "
        "relationnel et des jointures SQL."
    )
    doc.add_paragraph("Le modèle relationnel comprend :")

    add_styled_table(doc,
        ["Table", "Rôle", "Clés étrangères"],
        [
            ["PROMOTION", "Cohorte de diplômés (nom, année, filière)", "—"],
            ["ETUDIANT", "Profil alumni (15 champs dont compétences JSONB)", "id_promotion → PROMOTION"],
            ["ENTREPRISE", "Entreprise (nom, secteur, pays, ville)", "—"],
            ["EXPERIENCE_PRO", "Expérience professionnelle (poste, contrat, salaire)", "id_etudiant → ETUDIANT, id_entreprise → ENTREPRISE"],
            ["CERTIFICATION", "Certification professionnelle", "—"],
            ["OBTIENT", "Table d'association étudiant-certification", "id_etudiant → ETUDIANT, id_certification → CERTIFICATION"],
            ["CONSENTEMENT_RGPD", "Enregistrement de consentement", "id_etudiant → ETUDIANT"],
            ["QUESTIONNAIRE", "Questionnaire annuel", "—"],
            ["QUESTION", "Question d'un questionnaire", "id_questionnaire → QUESTIONNAIRE (CASCADE)"],
            ["REPONSE_QUESTIONNAIRE", "Réponse d'un alumni à un questionnaire", "id_etudiant → ETUDIANT, id_questionnaire → QUESTIONNAIRE"],
            ["AUDIT_LOG", "Journal des opérations de nettoyage", "—"],
        ]
    )

    doc.add_heading("6.2 Gouvernance des données", level=2)
    doc.add_paragraph(
        "La gouvernance des données repose sur trois piliers :"
    )

    doc.add_paragraph("Pilier 1 — Collecte :", style="List Bullet")
    doc.add_paragraph(
        "Deux méthodes d'import : /upload-etudiants/ (CSV/Excel via Pandas) et "
        "/import/excel (Excel via openpyxl). Chaque ligne est validée via Pydantic "
        "avant insertion. Les entreprises existantes sont préchargées en une seule "
        "requête pour optimiser les performances (pattern anti-N+1)."
    )

    doc.add_paragraph("Pilier 2 — Exploitation :", style="List Bullet")
    doc.add_paragraph(
        "Les endpoints admin calculent des indicateurs agrégés (taux d'emploi, "
        "salaire moyen, taux à 6 mois). Le filtrage avancé permet de croiser les "
        "données par promotion, secteur et entreprise. Les questions de questionnaire "
        "ont un tag pour identifier les KPI."
    )

    doc.add_paragraph("Pilier 3 — Archivage :", style="List Bullet")
    doc.add_paragraph(
        "Le workflow RGPD : détection des refus de consentement → anonymisation "
        "logique (UPDATE avec masquage) → journalisation dans AUDIT_LOG. "
        "L'archivage est non-destructif : les données sont masquées mais les "
        "enregistrements restent préservés pour l'intégrité référentielle."
    )

    doc.add_heading("6.3 Processus de mise à jour des données", level=2)
    doc.add_paragraph(
        "Le processus de mise à jour envisagé comprend :"
    )
    update_processes = [
        "Import Excel annuel : l'administrateur téléverse un fichier Excel contenant les mises à jour des profils alumni",
        "Questionnaire annuel automatisé : création d'un questionnaire actif via l'interface admin, les alumni y répondent via l'API /questionnaires/{id}/repondre",
        "Mise à jour individuelle : les alumni peuvent modifier leur profil et ajouter des expériences/certifications via les endpoints dédiés",
        "Nettoyage programmé : l'administrateur lance périodiquement les endpoints /admin/cleanup/orphelins et /admin/cleanup/doublons en mode dry-run d'abord, puis en mode exécution",
        "Archivage RGPD : exécuté après chaque campagne de consentement pour anonymiser les refus",
    ]
    for process in update_processes:
        doc.add_paragraph(process, style="List Bullet")

    doc.add_heading("6.4 Méthode de gestion de projet", level=2)
    doc.add_paragraph(
        "L'approche de gestion de projet est itérative, visible à travers :"
    )
    pm_items = [
        "Migrations numérotées (001→006) : chaque migration correspond à un incrément fonctionnel",
        "Architecture modulaire : le fichier unique initial (~590 lignes) a été découpé en 12 routers domain-specific",
        "Documentation continue : le README est mis à jour à chaque itération avec les corrections apportées",
        "Tests de robustesse : le README documente les corrections de sécurité (TOCTOU, messages d'erreur, variables d'environnement)",
        "Approche pragmatique : le pool de connexions est artisanal (documenté comme tel), l'auth admin est simplifiée (documentée comme nécessitant une évolution vers OAuth2)",
    ]
    for item in pm_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("6.5 Limites connues et évolutions envisagées", level=2)
    doc.add_paragraph(
        "Le README documente explicitement les limites du projet :"
    )
    limitations = [
        "L'authentification admin (X-API-Key) est volontairement simple : si l'application doit distinguer plusieurs profils utilisateurs (école vs alumni), il faudra passer à OAuth2/JWT",
        "Le pool de connexions est artisanal (file d'attente simple) ; en production, un vrai pool (SQLAlchemy QueuePool, ou PgBouncer) gérerait la détection des connexions mortes",
        "L'import de fichier reste ligne par ligne pour les étudiants (nécessaire pour récupérer chaque id_etudiant) ; un vrai import de masse (PostgreSQL COPY) serait à envisager pour les gros volumes",
    ]
    for lim in limitations:
        doc.add_paragraph(lim, style="List Bullet")

    # ═══════════════════════════════════════════════════════════
    # ANNEXE : ENDPOINTS API
    # ═══════════════════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading("Annexe : Liste des endpoints API", level=1)

    endpoints_data = [
        ["GET", "/", "Message de bienvenue"],
        ["POST", "/promotions/", "Créer une promotion"],
        ["GET", "/promotions/", "Lister les promotions (paginé, filtrable)"],
        ["GET", "/promotions/{id}", "Détail d'une promotion"],
        ["DELETE", "/promotions/{id}", "Supprimer une promotion (+ cascade)"],
        ["POST", "/etudiants/", "Créer un étudiant"],
        ["GET", "/etudiants/", "Lister les étudiants (paginé, recherchable)"],
        ["GET", "/etudiants/{id}", "Détail d'un étudiant (avec promo, entreprise, nb exp)"],
        ["PUT", "/etudiants/{id}", "Mise à jour complète d'un étudiant"],
        ["PATCH", "/etudiants/{id}", "Mise à jour partielle d'un étudiant"],
        ["DELETE", "/etudiants/{id}", "Supprimer un étudiant (+ cascade)"],
        ["POST", "/entreprises/", "Créer une entreprise"],
        ["GET", "/entreprises/", "Lister les entreprises (paginé, filtrable)"],
        ["GET", "/entreprises/{id}", "Détail d'une entreprise"],
        ["PUT", "/entreprises/{id}", "Mettre à jour une entreprise"],
        ["DELETE", "/entreprises/{id}", "Supprimer une entreprise"],
        ["GET", "/etudiants/{id}/experiences", "Expériences d'un étudiant"],
        ["GET", "/experiences/", "Rechercher des expériences (paginé, filtrable)"],
        ["POST", "/experiences/", "Créer une expérience"],
        ["DELETE", "/experiences/{id}", "Supprimer une expérience"],
        ["POST", "/etudiants/{id}/experiences", "Alumni ajoute une expérience"],
        ["POST", "/certifications/", "Créer une certification"],
        ["GET", "/certifications/", "Lister les certifications"],
        ["DELETE", "/certifications/{id}", "Supprimer une certification"],
        ["POST", "/etudiants-certifications/", "Associer certification à un étudiant"],
        ["DELETE", "/etudiants-certifications/", "Dissocier certification d'un étudiant"],
        ["GET", "/consentements/etudiant/{id}", "Historique des consentements RGPD"],
        ["POST", "/consentements/", "Enregistrer un consentement"],
        ["DELETE", "/consentements/{id}", "Supprimer un consentement"],
        ["GET", "/admin/etudiants/filtrer", "Filtrage avancé des alumni"],
        ["GET", "/admin/indicateurs", "Indicateurs d'insertion (taux emploi, salaire...)"],
        ["GET", "/admin/indicateurs/secteurs", "Répartition par secteur"],
        ["GET", "/admin/indicateurs/kpi-tag", "KPI par tag de question"],
        ["GET", "/admin/cleanup/orphelins", "Détection des orphelins (dry-run)"],
        ["DELETE", "/admin/cleanup/orphelins", "Suppression des orphelins"],
        ["GET", "/admin/cleanup/doublons", "Détection des doublons (dry-run)"],
        ["DELETE", "/admin/cleanup/doublons", "Suppression des doublons"],
        ["POST", "/admin/cleanup/rgpd/archiver", "Anonymisation RGPD"],
        ["GET", "/admin/cleanup/audit", "Historique des audits"],
        ["POST", "/admin/questionnaires/", "Créer un questionnaire"],
        ["GET", "/admin/questionnaires/", "Lister les questionnaires"],
        ["GET", "/admin/questionnaires/{id}", "Détail d'un questionnaire"],
        ["PUT", "/admin/questionnaires/{id}", "Modifier un questionnaire"],
        ["DELETE", "/admin/questionnaires/{id}", "Supprimer un questionnaire"],
        ["PATCH", "/admin/questionnaires/{id}/desactiver", "Désactiver un questionnaire"],
        ["PATCH", "/admin/questionnaires/{id}/reactiver", "Réactiver un questionnaire"],
        ["GET", "/admin/questionnaires/{id}/reponses", "Voir les réponses"],
        ["GET", "/questionnaires/actif", "Questionnaire actif (alumni)"],
        ["GET", "/questionnaires/etudiant/{id}/reponses", "Mes réponses (alumni)"],
        ["POST", "/questionnaires/{id}/repondre", "Répondre au questionnaire (alumni)"],
        ["POST", "/import/excel", "Importer des alumni depuis Excel"],
        ["GET", "/import/template", "Télécharger le template d'import"],
        ["GET", "/import/export/alumni", "Exporter tous les alumni en Excel"],
        ["POST", "/upload-etudiants/", "Import CSV/Excel (automatisation)"],
        ["POST", "/auth/otp/request", "Demander un code OTP"],
        ["POST", "/auth/otp/verify", "Vérifier le code OTP → JWT"],
        ["POST", "/auth/admin/login", "Connexion admin → JWT"],
    ]

    add_styled_table(doc,
        ["Méthode", "Chemin", "Description"],
        endpoints_data
    )

    # ── Sauvegarde ──
    output_path = r"C:\Users\PC\OneDrive\Desktop\stage\alumni_crm_api\Preparation_Oral_Stage_CRM_Alumni.docx"
    doc.save(output_path)
    print(f"Document généré avec succès : {output_path}")


if __name__ == "__main__":
    generer_document()
