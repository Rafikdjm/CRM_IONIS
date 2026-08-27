"""Génère le rapport d'audit Word (Audit Final Alumni CRM)."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Styles globaux ────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

# ── Helper : couleur de fond d'en-tête de tableau ────────────────
def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_status_cell(cell, status):
    color_map = {
        "FAIT": "27AE60",
        "PARTIEL": "F39C12",
        "MANQUANT": "E74C3C",
    }
    cell.text = status
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(color_map.get(status, "000000"))
            run.font.size = Pt(10)


def make_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # en-tête
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "1B3A5C")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
    # données
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if headers[c_idx] == "Statut":
                add_status_cell(cell, val)
            else:
                cell.text = str(val)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(10)
    return table


# ═══════════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ═══════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AUDIT FINAL & COMPLET")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Projet Alumni CRM")
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("par rapport au cahier des charges\ndu sujet de stage PreMsc 2026")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph("")

date_audit = datetime.date.today().strftime("%d/%m/%Y")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"Date de l'audit : {date_audit} (mise a jour du document initial du 19/08/2026)\nÉchéance soutenance : 19 septembre 2026")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# TABLE DES MATIÈRES (texte)
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("Table des matières", level=1)
toc_items = [
    "1. MANAGEMENT",
    "   1.1 Cartographie des données",
    "   1.2 Charte RGPD / consentement",
    "   1.3 Stratégie de mise à jour des données",
    "   1.4 Indicateurs d'insertion",
    "2. INFORMATIQUE",
    "   2.1 Modélisation de la base de données",
    "   2.2 Interface admin",
    "   2.3 Interface étudiant / alumni",
    "   2.4 Import / Export",
    "3. LIVRABLES ATTENDUS",
    "   3.1 Rapport de stage",
    "   3.2 MCD / MLD",
    "   3.3 Prototype fonctionnel",
    "   3.4 Guide des processus d'animation du réseau",
    "4. TABLEAU RÉCAPITULATIF",
    "5. PLAN D'ACTION PRIORITAIRE AVANT LA SOUTENANCE",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 1. MANAGEMENT
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("1. MANAGEMENT", level=1)

# ── 1.1 Cartographie des données ─────────────────────────────────
doc.add_heading("1.1 Cartographie des données", level=2)
p = doc.add_paragraph()
run = p.add_run("Statut : FAIT")
run.bold = True
run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

doc.add_heading("Données d'entrée (collectées via l'interface)", level=3)
doc.add_paragraph(
    "• Coordonnées : address, city, country, telephone, linkedin, email, "
    "email_academique, date_naissance — collectés à l'inscription "
    "(AlumniRegistration.jsx) et modifiables via le profil "
    "(AlumniProfile.jsx)",
    style="List Bullet",
)
doc.add_paragraph(
    "• Parcours antérieur : parcours_anterieur (text) — collecté à l'inscription "
    "(champs previous_education + previous_school dans AlumniRegistration.jsx:18-19) "
    "et affiché/modifiable dans AlumniProfile.jsx:257 et AlumniEditModal.jsx:236",
    style="List Bullet",
)
doc.add_paragraph(
    "• Compétences : skills (JSONB) — géré via tags dans AlumniProfile.jsx",
    style="List Bullet",
)

doc.add_heading("Données de sortie", level=3)
doc.add_paragraph(
    "• Entreprise : table ENTREPRISE (nom, secteur, pays, ville) liée via EXPERIENCE_PRO",
    style="List Bullet",
)
doc.add_paragraph(
    "• Poste : intitule_poste, type_contrat, poste_actuel dans EXPERIENCE_PRO",
    style="List Bullet",
)
doc.add_paragraph(
    "• Salaire : colonnes salaire (numeric) et salary_annuel NUMERIC(10,2) "
    "(migration 012_salary_annuel.sql) dans EXPERIENCE_PRO — saisie via un select "
    "de tranches chiffrées SALARY_RANGES (AlumniCareer.jsx:9-22, select en "
    "AlumniCareer.jsx:151-156) converti en salary_annuel côté client "
    "(api.js:336-343), exploité par le dashboard (admin.py:145 calculer_indicateurs)",
    style="List Bullet",
)

doc.add_heading("Modélisation", level=3)
doc.add_paragraph(
    "Toutes ces colonnes existent dans erd_alumni_crm.mmd (lignes 34-70). "
    "La relation ENTREPRISE → EXPERIENCE_PRO est de type 1-N confirmée."
)

# ── 1.2 Charte RGPD ──────────────────────────────────────────────
doc.add_heading("1.2 Charte RGPD / consentement", level=2)
p = doc.add_paragraph()
run = p.add_run("Statut : FAIT")
run.bold = True
run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

rgpd_items = [
    ("Consentement horodaté",
     "CONSENTEMENT_RGPD.date_consentement (DATE NOT NULL) — erd_alumni_crm.mmd:84"),
    ("Modifiable",
     "UPSERT sur (id_etudiant, type_consentement) UNIQUE — migration 006_consentement_upsert.sql, endpoint rgpd.py:41"),
    ("Traçé en base",
     "Chaque action loguée dans AUDIT_LOG avec acteur (admin:nom | alumni:id | system) — erd_alumni_crm.mmd:119"),
    ("Export fonctionnel",
     "demandes_rgpd.py — workflow complet : demande → traitement → export JSON / anonymisation avec CASCADE (purge.py)"),
    ("Suppression fonctionnelle",
     "Anonymisation différée (6 mois par défaut via PURGE_DELAY_MONTHS) avec date_anonymisation sur ETUDIANT, FK CASCADE sur REPONSE_QUESTIONNAIRE — migration 008_purge_anonymises.sql"),
    ("Interface alumni",
     "AlumniConsent.jsx — toggle avec sauvegarde serveur"),
    ("Interface admin",
     "AdminRgpdDemandes.jsx (618 lignes) — workflow complet avec prise en charge, traitement, rejet, export, suppression"),
]
for title, desc in rgpd_items:
    doc.add_paragraph(f"• {title} : {desc}", style="List Bullet")

# ── 1.3 Stratégie de mise à jour ─────────────────────────────────
doc.add_heading("1.3 Stratégie de mise à jour des données (gouvernance)", level=2)
p = doc.add_paragraph()
run = p.add_run("Statut : FAIT")
run.bold = True
run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

doc.add_heading("Questionnaire annuel — FAIT", level=3)
doc.add_paragraph(
    "Le système de questionnaire annuel existe et fonctionne :"
)
doc.add_paragraph("Tables : QUESTIONNAIRE, QUESTION, REPONSE_QUESTIONNAIRE — migration 003_questionnaire_annuel.sql", style="List Bullet")
doc.add_paragraph("Création admin : AdminQuestionnaires.jsx + questionnaires.py", style="List Bullet")
doc.add_paragraph("Réponses alumni : AlumniSurvey.jsx", style="List Bullet")
doc.add_paragraph("KPI via tag sur QUESTION (migration 004_question_tag.sql) et conditionnee_statut_emploi (migration 005_question_statut_emploi.sql)", style="List Bullet")
doc.add_paragraph("Indicateurs calculés dans admin.py (calculer_indicateurs, admin.py:145) : taux d'emploi, salaire moyen, taux 6 mois, taux adéquation", style="List Bullet")

doc.add_heading("Relance proactive — FAIT", level=3)
p = doc.add_paragraph()
run = p.add_run(
    "L'endpoint POST /admin/questionnaires/notififier (questionnaires.py:456) envoie "
    "une notification email aux alumni n'ayant pas encore répondu au questionnaire "
    "actif, avec filtre optionnel par promotion. L'envoi passe par Resend "
    "(comme les OTP), le corps du message est générique et ne contient pas encore "
    "de lien direct vers le formulaire."
)
run.italic = True

doc.add_paragraph("Reste à faire (non bloquant, suites à donner au projet) :", style="List Bullet")
doc.add_paragraph("Un déclenchement automatique planifié (cron hebdomadaire) en plus de l'appel manuel de l'endpoint", style="List Bullet 2")
doc.add_paragraph("Une interface admin dédiée au déclenchement de la relance (aujourd'hui appel API uniquement)", style="List Bullet 2")
doc.add_paragraph("Idéalement 2-3 relances espacées, puis un flag \"injoignable\"", style="List Bullet 2")

# ── 1.4 Indicateurs d'insertion ──────────────────────────────────
doc.add_heading("1.4 Indicateurs d'insertion", level=2)
p = doc.add_paragraph()
run = p.add_run("Statut : FAIT")
run.bold = True
run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

indicators = [
    ("Taux d'emploi global", "admin.py:145 (calculer_indicateurs)", "Hero card AdminDashboard.jsx:92"),
    ("Taux d'emploi à 6 mois", "admin.py:145 (calculer_indicateurs)", "KPI card AdminDashboard.jsx"),
    ("Taux d'adéquation formation/emploi", "admin.py:692 (/indicateurs/kpi-tag)", "KPI card"),
    ("Salaire moyen", "admin.py:145 (calculer_indicateurs)", "Jauge demi-cercle dynamique AdminDashboard.jsx"),
    ("Taux de recommandation", "admin.py (calculer_indicateurs)", "KPI card"),
    ("Répartition sectorielle", "admin.py:402 (/indicateurs/secteurs)", "Donut chart"),
    ("Évolution par promotion", "admin.py (calculer_indicateurs)", "Bar chart"),
    ("Taux de complétion profil", "admin.py (calculer_indicateurs)", "KPI card"),
    ("Types de contrat", "admin.py:452 (/indicateurs/types-contrat)", "Bar chart horizontal"),
    ("Maturité des cohortes", "admin.py (taux_6_mois par promo)", "Timeline"),
]
make_table(doc, ["Indicateur", "API (admin.py)", "Dashboard (AdminDashboard.jsx)"], indicators)

doc.add_paragraph("")
doc.add_paragraph(
    "Confirmé : Le \"taux d'emploi à 6 mois\" et le \"taux d'adéquation formation/emploi\" "
    "sont explicitement calculés et affichés, comme requis par le sujet de stage."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 2. INFORMATIQUE
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("2. INFORMATIQUE", level=1)

# ── 2.1 Modélisation BDD ─────────────────────────────────────────
doc.add_heading("2.1 Modélisation de la base de données", level=2)
p = doc.add_paragraph()
run = p.add_run("Statut : FAIT")
run.bold = True
run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

doc.add_heading("MCD/MLD existant", level=3)
mcd_files = [
    ("erd_alumni_crm.mmd", "14 tables", "COMPLET et à jour (généré 10/08/2026, vérifié 15/08, 22/08 et re-audité 27/08/2026)", "FAIT"),
    ("erd_alumni_crm.docx", "Word", "Version Word régénérée depuis le .mmd (27/08/2026)", "FAIT"),
    ("000_schema_initial.sql + migrations 001-015", "16 fichiers SQL", "Rejeu complet validé sur base vide : les 14 tables sont recréées", "FAIT"),
    ("MCD_MLD V2.loo", "Binaire Looping", "Phase initiale de conception (28/07/2026), supplanté par Mermaid ; les autres fichiers Looping (.loo/.lo1 doublons) ont été supprimés", "Source historique"),
]
make_table(doc, ["Fichier", "Description", "Statut", "Verdict"], mcd_files)

doc.add_paragraph("")
doc.add_heading("Relation 1-N étudiant → expériences", level=3)
doc.add_paragraph(
    "CONFIRMÉE — erd_alumni_crm.mmd:18 : ETUDIANT ||--o{ EXPERIENCE_PRO (0,n expériences par étudiant). "
    "Historique de carrière multiple, pas un seul poste."
)

# ── 2.2 Interface admin ──────────────────────────────────────────
doc.add_heading("2.2 Interface admin", level=2)
p = doc.add_paragraph()
run = p.add_run("Statut : FAIT")
run.bold = True
run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

doc.add_heading("Filtrage combiné promotion + secteur + entreprise", level=3)
doc.add_paragraph(
    "CONFIRMÉ — AlumniDirectory.jsx (686 lignes) : filtres AND-combinés dans un useMemo (l.72-114) :"
)
filter_items = [
    "Promotion : serveur (params.promotion, l.43)",
    "Secteur : client (a.sector, l.76-79), incluant \"Autre\"",
    "Entreprise : client (a.current_company, l.86-88)",
    "Disponibilité, compétence, contact autorisé, anonymisation : aussi disponibles",
    "Recherche texte sur nom/email/entreprise (l.103-111)",
]
for f in filter_items:
    doc.add_paragraph(f, style="List Bullet")

doc.add_heading("Dashboard avec indicateurs d'insertion", level=3)
doc.add_paragraph(
    "CONFIRMÉ — AdminDashboard.jsx (1208 lignes) : hero cards, jauge salaire dynamique, "
    "donuts, bar charts, timeline cohortes, KPI tags depuis questionnaires."
)

# ── 2.3 Interface alumni ─────────────────────────────────────────
doc.add_heading("2.3 Interface étudiant / alumni", level=2)
p = doc.add_paragraph()
run = p.add_run("Statut : FAIT")
run.bold = True
run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

alumni_checks = [
    ("Formulaire d'inscription INITIAL distinct",
     "CONFIRMÉ — AlumniRegistration.jsx (601 lignes) : wizard multi-étapes (4 étapes), "
     "distinct de la mise à jour. Stocke alumni_id en localStorage après succès."),
    ("Mise à jour du profil professionnel",
     "CONFIRMÉE — AlumniProfile.jsx : consultation et mise à jour du profil ; "
     "AlumniCareer.jsx (671 lignes) : CRUD expériences + certifications avec modals de confirmation. "
     "Limite assumée : pas de modification in place d'une expérience existante (ajout/suppression uniquement)."),
    ("Suppression d'une entrée par erreur",
     "CONFIRMÉE — AlumniCareer.jsx : bouton \"Supprimer\" + modal de confirmation. "
     "confirmRemove() (l.403) appelle careerAPI.delete() côté serveur. "
     "Bloqué sur comptes anonymisés."),
]
for title, desc in alumni_checks:
    doc.add_paragraph(f"• {title} : {desc}", style="List Bullet")

# ── 2.4 Import / Export ──────────────────────────────────────────
doc.add_heading("2.4 Import / Export", level=2)
p = doc.add_paragraph()
run = p.add_run("Statut : FAIT")
run.bold = True
run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

doc.add_paragraph("Import Excel — CONFIRMÉ :")
import_items = [
    "import_export.py:186 : POST /import/excel — accepte .xlsx, .xls, .csv",
    "25 colonnes supportées (COLUMN_MAP, import_export.py:28) : prénom, nom, email, téléphone, promotion, entreprise, poste, secteur, salaire, etc.",
    "Logique complète : création ETUDIANT + ENTREPRISE + EXPERIENCE_PRO, résolution de doublons email",
    "Template de téléchargement (GET /import/template, import_export.py:350)",
    "Frontend : ExcelImport.jsx (366 lignes) — drag & drop, preview 10 lignes, reconnaissance de colonnes",
]
for item in import_items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph("")
doc.add_paragraph("Export — CONFIRMÉ :")
doc.add_paragraph("GET /import/export/alumni (import_export.py:383) — génère .xlsx", style="List Bullet")
doc.add_paragraph("Fallback client-side dans ExcelImport.jsx:119-156", style="List Bullet")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 3. LIVRABLES ATTENDUS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("3. LIVRABLES ATTENDUS", level=1)

# ── 3.1 Rapport de stage ─────────────────────────────────────────
doc.add_heading("3.1 Rapport de stage", level=2)
p = doc.add_paragraph()
run = p.add_run("Statut : FAIT")
run.bold = True
run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

doc.add_paragraph(
    "Le rapport de stage existe : source Markdown (Rapport/rapport.md) et PDF "
    "généré par script (Rapport/generate_reports.py)."
)
rapport_items = [
    "Structure complète : résumé/abstract, contexte IONIS-STM, missions, bilan de compétences, difficultés/solutions, bibliographie, annexes",
    "Couvre les 5 domaines de mission du sujet (modélisation, backend, frontend, RGPD, indicateurs)",
    "Quelques compléments rédactionnels restent à porter par l'auteur (tuteur, dates, captures des annexes)",
]
for item in rapport_items:
    doc.add_paragraph(item, style="List Bullet")

# ── 3.2 MCD/MLD ─────────────────────────────────────────────────
doc.add_heading("3.2 MCD / MLD", level=2)
p = doc.add_paragraph()
run = p.add_run("Statut : FAIT (avec réserve)")
run.bold = True
run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

mcd_livrable = [
    ("Fichier principal", "alumni_crm_api/docs/erd_alumni_crm.mmd — 14 tables, re-audité 27/08/2026"),
    ("Version Word", "alumni_crm_api/docs/erd_alumni_crm.docx — régénéré depuis le .mmd (27/08/2026)"),
    ("Fichier Looping", "MCD_MLD V2.loo (28/07/2026) — phase initiale, supplanté par Mermaid"),
    ("Rejeu base vide", "000_schema_initial.sql + migrations 000-015 : les 16 migrations reconstruisent intégralement le schéma"),
]
for title, desc in mcd_livrable:
    doc.add_paragraph(f"• {title} : {desc}", style="List Bullet")

# ── 3.3 Prototype fonctionnel ────────────────────────────────────
doc.add_heading("3.3 Prototype fonctionnel", level=2)
p = doc.add_paragraph()
run = p.add_run("Statut : FAIT")
run.bold = True
run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

proto_items = [
    "Backend FastAPI opérationnel (16 routeurs montés, 85 endpoints)",
    "Frontend React/Vite complet (dist/ présent avec build production)",
    "Base PostgreSQL (14 tables, 16 migrations)",
    "Auth OTP + JWT (alumni), API Key + JWT (admin)",
    "Tests : aucune suite de tests automatisés conservée dans le dépôt à l'issue du stage (les tests pytest backend ont été retirés) ; validation par scripts ad hoc, rejeu des migrations et tests manuels",
    "Import/Export Excel, RGPD bout en bout, questionnaire annuel",
]
for item in proto_items:
    doc.add_paragraph(item, style="List Bullet")

# ── 3.4 Guide animation réseau ───────────────────────────────────
doc.add_heading("3.4 Guide des processus d'animation du réseau", level=2)
p = doc.add_paragraph()
run = p.add_run("Statut : FAIT")
run.bold = True
run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

doc.add_paragraph(
    "Guide standalone dédié au service des relations entreprises, généré par "
    "script : \"Guide des Processus - Animation du Reseau Alumni.pdf\" "
    "(Rapport/generate_reports.py)."
)
animation_items = [
    "Pilotage des campagnes de questionnaire (création, relances via POST /admin/questionnaires/notififier)",
    "Valorisation du réseau (filtrage dashboard, identification partenariats)",
    "Newsletter (levier d'animation, ciblage par consentement via POST /newsletter/envoyer)",
]
for item in animation_items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 4. TABLEAU RÉCAPITULATIF
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("4. TABLEAU RÉCAPITULATIF", level=1)

recap = [
    ("1", "Cartographie des données", "FAIT", "AlumniRegistration.jsx, AlumniProfile.jsx, AlumniCareer.jsx, erd_alumni_crm.mmd"),
    ("2", "Charte RGPD / consentement", "FAIT", "rgpd.py, demandes_rgpd.py, purge.py, AlumniConsent.jsx, AdminRgpdDemandes.jsx, migrations 006-010"),
    ("3", "Stratégie de mise à jour", "FAIT", "questionnaires.py (relance : POST /admin/questionnaires/notififier), AdminQuestionnaires.jsx, AlumniSurvey.jsx — restent cron + UI dédiée (non bloquant)"),
    ("4", "Indicateurs d'insertion", "FAIT", "admin.py:145+ (/indicateurs...), AdminDashboard.jsx"),
    ("5", "Modélisation BDD", "FAIT", "erd_alumni_crm.mmd (14 tables) ; mcd_corrige.md supprimé"),
    ("6", "Interface admin", "FAIT", "AlumniDirectory.jsx:72-114, AdminDashboard.jsx, AdminPromotions.jsx"),
    ("7", "Interface alumni", "FAIT", "AlumniRegistration.jsx, AlumniProfile.jsx, AlumniCareer.jsx"),
    ("8", "Import / Export", "FAIT", "import_export.py (COLUMN_MAP 25 colonnes), ExcelImport.jsx"),
    ("9", "Rapport de stage", "FAIT", "Rapport/rapport.md + Rapport de Stage - Alumni CRM.pdf (compléments rédactionnels en cours)"),
    ("10", "MCD / MLD", "FAIT", "erd_alumni_crm.mmd, erd_alumni_crm.docx, MCD_MLD V2.loo"),
    ("11", "Prototype fonctionnel", "FAIT", "dist/, requirements.txt, 16 routeurs montés (14 fichiers), 85 endpoints"),
    ("12", "Guide animation réseau", "FAIT", "Guide des Processus - Animation du Reseau Alumni.pdf (standalone)"),
]
make_table(doc, ["#", "Point", "Statut", "Fichiers / Routes concernés"], recap)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 5. PLAN D'ACTION
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("5. PLAN D'ACTION PRIORITAIRE AVANT LA SOUTENANCE (19 septembre 2026)", level=1)

doc.add_heading("Urgence CRITIQUE — Manques documentaires (rédaction)", level=2)
critical = [
    ("1", "Rédiger le rapport de stage",
     "FAIT depuis la version initiale de l'audit : Rapport/rapport.md + PDF généré.",
     "—",
     "Restent les compléments rédactionnels portés par l'auteur (tuteur, dates, captures des annexes)"),
    ("2", "Rédiger le guide d'animation du réseau",
     "FAIT depuis la version initiale de l'audit : Guide des Processus - Animation du Reseau Alumni.pdf.",
     "—",
     "Document standalone généré par Rapport/generate_reports.py"),
]
make_table(doc, ["#", "Action", "Justification", "Estimation", "Détails"], critical)

doc.add_paragraph("")
doc.add_heading("Urgence MOYENNE — Manques techniques (code)", level=2)
medium = [
    ("3", "Compléter la relance proactive",
     "PARTIELLEMENT FAIT : POST /admin/questionnaires/notififier envoie déjà les relances email aux non-répondants (filtre promotion).",
     "1-2 jours",
     "Reste : déclenchement planifié (cron hebdomadaire), interface admin dédiée au déclenchement, 2-3 relances espacées + flag \"injoignable\""),
    ("4", "Synchroniser mcd_corrige.md",
     "FAIT : le fichier obsolète a été supprimé du dépôt ; erd_alumni_crm.mmd fait foi (14 tables).",
     "—",
     "Rien à faire"),
]
make_table(doc, ["#", "Action", "Justification", "Estimation", "Détails"], medium)

doc.add_paragraph("")
doc.add_heading("Attention — Non bloquant", level=2)
doc.add_paragraph("Rejouabilité base vide rétablie : migration 000_schema_initial.sql (bootstrap des tables métier) + 013_otp_codes.sql (table OTP, anciennement créée hors migrations) — rejeu complet 000→015 (16 migrations) validé sur base vide, 14/14 tables recréées", style="List Bullet")
doc.add_paragraph("Les fichiers Looping obsolètes (MCD_MLD.loo, MCD_MLD.lo1) ont été supprimés — seul MCD_MLD V2.loo est conservé comme source", style="List Bullet")
doc.add_paragraph("Suites techniques documentées dans le README API : reconstituer le script E2E (la démarche est décrite), introduire des tests backend automatisés (pytest)", style="List Bullet")

# ── Sauvegarde ────────────────────────────────────────────────────
output_path = r"C:\Users\PC\OneDrive\Desktop\stage\Audit_Final_Alumni_CRM.docx"
doc.save(output_path)
print(f"Document généré : {output_path}")
