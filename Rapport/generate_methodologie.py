import os
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

HEADER_FILL = 'DBEAFE'      # bleu clair
STRIPE_FILL = 'F4F7FC'      # gris très clair (alternance)
BODY_FILL = 'FFFFFF'
BORDER_HEX = 'C4D2E2'       # bordure fine gris-bleu
HEADER_TEXT = RGBColor(0x1E, 0x29, 0x3B)
BODY_TEXT = RGBColor(0x37, 0x41, 0x51)
TABLE_FONT_PT = 10


def set_cell_bg(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def _set_cell_borders(cell, hexcolor=BORDER_HEX, width='4'):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), width)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), hexcolor)
        borders.append(el)
    tcPr.append(borders)


def _write_cell(cell, text, bold=False, color=BODY_TEXT, size=TABLE_FONT_PT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.font.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color


def _set_row_keep_together(row):
    """w:cantSplit → une ligne ne peut pas être coupée entre deux pages."""
    trPr = row._tr.get_or_add_trPr()
    cant = OxmlElement('w:cantSplit')
    trPr.append(cant)


def _set_row_repeat_header(row):
    """w:tblHeader → l'en-tête est répété quand le tableau se poursuit sur une page."""
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement('w:tblHeader')
    trPr.append(th)


def _set_col_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, w in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(w)


def styled_table(doc, headers, rows, widths_cm, header_bold_cols=None):
    """Table uniforme : en-tête bleu clair gras, bordures fines, alternance de lignes."""
    table = doc.add_table(rows=0, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr = table.add_row().cells
    _set_row_repeat_header(table.rows[0])
    _set_row_keep_together(table.rows[0])
    for i, h in enumerate(headers):
        set_cell_bg(hdr[i], HEADER_FILL)
        _set_cell_borders(hdr[i])
        _write_cell(hdr[i], h, bold=True, color=HEADER_TEXT)

    header_bold_cols = header_bold_cols or set()
    for r, row in enumerate(rows):
        cells = table.add_row().cells
        _set_row_keep_together(table.rows[r + 1])
        fill = STRIPE_FILL if r % 2 == 1 else BODY_FILL
        for i, val in enumerate(row):
            set_cell_bg(cells[i], fill)
            _set_cell_borders(cells[i])
            _write_cell(cells[i], val, bold=(i in header_bold_cols))

    _set_col_widths(table, widths_cm)
    doc.add_paragraph()
    return table


def titre(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    return p


def sous_titre(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    return p


def chapitre(doc, num, text):
    p = doc.add_paragraph()
    run = p.add_run(f"{num}. {text}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    return p


def section(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    return p


def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    return p


def indicateur_table(doc, rows):
    return styled_table(
        doc,
        headers=['Champ', 'Valeur'],
        rows=[(champ, valeur) for champ, valeur in rows],
        widths_cm=[4.0, 12.0],
        header_bold_cols={0},
    )


def generate():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)
    sec.top_margin = Cm(2.4)
    sec.bottom_margin = Cm(2.6)
    titre(doc, "Methodologie des indicateurs d'insertion")
    sous_titre(doc, "CRM Alumni (IONIS STM) - Fiche de reference complete et fidele au code")
    body(doc, "Chaque indicateur est documente avec sa formule exacte, ses tables/colonnes, ses exclusions et un exemple chiffre calcule comme le font les requetes SQL reelles du backend (routers/admin.py).")

    # 1. Principes
    chapitre(doc, 1, "Principes generaux")
    bullet(doc, "Comptes anonymises (ETUDIANT.date_anonymisation IS NOT NULL) exclus de tous les indicateurs.")
    bullet(doc, "Source de verite de l'emploi : EXPERIENCE_PRO (poste_actuel / experience en cours), pas availability_status.")
    bullet(doc, "Salaire : salary_annuel prioritaire si renseigne, sinon repli sur le champ texte salaire.")
    bullet(doc, "Une date de fin deja passee exclut toujours l'experience, meme si poste_actuel = TRUE.")

    # 2. Indicateurs
    chapitre(doc, 2, "Les 10 indicateurs avec exemples chiffres")

    section(doc, "2.1 Total Alumni actifs")
    indicateur_table(doc, [
        ("Formule", "COUNT(*) FROM ETUDIANT WHERE date_anonymisation IS NULL"),
        ("Tables / colonnes", "ETUDIANT.date_anonymisation"),
        ("Exclusions / limites", "Comptes anonymises exclus."),
        ("Exemple chiffre", "3 alumni (Rafik, Alice, Anonyme) -> 2 (Anonyme exclu)"),
    ])

    section(doc, "2.2 Alumni actifs (>= 1 experience)")
    indicateur_table(doc, [
        ("Formule", "COUNT(DISTINCT id_etudiant) FROM EXPERIENCE_PRO JOIN ETUDIANT non anonymises"),
        ("Tables / colonnes", "EXPERIENCE_PRO.id_etudiant, ETUDIANT.date_anonymisation"),
        ("Exclusions / limites", "Doublons exclus (DISTINCT). Aucune condition de date."),
        ("Exemple chiffre", "Rafik (2 exp), Alice (1 exp), Anonyme exclu -> 2"),
    ])

    section(doc, "2.3 Taux de completion / couverture")
    indicateur_table(doc, [
        ("Formule", "(alumni avec >= 1 experience / total alumni) x 100"),
        ("Tables / colonnes", "EXPERIENCE_PRO.id_etudiant, ETUDIANT.date_anonymisation"),
        ("Exclusions / limites", "Profil sans experience = incomplet."),
        ("Exemple chiffre", "2 non anonymises, 1 avec exp -> (1/2)x100 = 50 %"),
    ])

    section(doc, "2.4 Taux d'emploi a 6 mois")
    indicateur_table(doc, [
        ("Formule", "date_reference = annee_diplome || '-12-01' (diplomation juin + 6 mois). Exp. active si date_debut <= reference AND (date_fin IS NULL OR date_fin >= reference). Taux = SUM(emplois_6_mois)/SUM(total_diplomes) x 100 (promotions matures)."),
        ("Tables / colonnes", "PROMOTION.annee_diplome, ETUDIANT.id_promotion, EXPERIENCE_PRO.date_debut/fin"),
        ("Exclusions / limites", "Promotions non matures exclues du global -> statut en_attente, taux null."),
        ("Exemple chiffre", "2024 : 7/10 ; 2025 : 9/12 ; 2026 en_attente -> (16/22) = 72,7 %"),
    ])

    section(doc, "2.5 Taux d'emploi global (brut)")
    indicateur_table(doc, [
        ("Formule", "(etudiants en poste / total etudiants) x 100, par promotion puis somme"),
        ("Tables / colonnes", "EXPERIENCE_PRO.poste_actuel, ETUDIANT.id_promotion"),
        ("Exclusions / limites", "Un seul comptage par etudiant (DISTINCT). availability_status non utilise."),
        ("Exemple chiffre", "Promo 2024 : 8 en poste / 10 -> 80 %"),
    ])

    section(doc, "2.6 Adequation formation/emploi")
    indicateur_table(doc, [
        ("Formule", "(reponses 'Oui' / reponses exploitables) x 100. Question retenue : tag = 'adequation_formation' + questionnaire actif le plus recent."),
        ("Tables / colonnes", "QUESTION.tag, QUESTIONNAIRE.actif, REPONSE_QUESTIONNAIRE.reponses (JSON)"),
        ("Exclusions / limites", "Reponses 'non applicable' exclues du denominateur. Si aucune question taguee -> etat vide."),
        ("Exemple chiffre", "5 reponses dont 1 'Non applicable' (Oui,Oui,Oui,Non) -> 3/4 exploitables = 75 %"),
    ])

    section(doc, "2.7 Repartition par secteur")
    indicateur_table(doc, [
        ("Formule", "Par secteur : COUNT(DISTINCT id_etudiant) ayant poste_actuel = TRUE. % = count / total alumni actifs x 100."),
        ("Tables / colonnes", "ENTREPRISE.secteur_activite, EXPERIENCE_PRO.poste_actuel"),
        ("Exclusions / limites", "Secteur vide exclu. Un alumni multi-postes peut apparaitre dans plusieurs secteurs."),
        ("Exemple chiffre", "Info 3 (50%), Finance 2 (33%), Sante 1 (17%) -> donut"),
    ])

    section(doc, "2.8 Alumni par promotion")
    indicateur_table(doc, [
        ("Formule", "Par promotion : effectif, % en poste, couverture experience, salaire moyen. Promotions sans alumni exclues (HAVING COUNT > 0)."),
        ("Tables / colonnes", "PROMOTION, ETUDIANT.id_promotion, EXPERIENCE_PRO.poste_actuel/salaire"),
        ("Exclusions / limites", "Salaire moyen sur postes actuels uniquement."),
        ("Exemple chiffre", "2024 : 10, 80% ; 2025 : 12, 75% ; 2026 : 15, 73%"),
    ])

    section(doc, "2.9 Salaire moyen")
    indicateur_table(doc, [
        ("Formule", "AVG(CASE WHEN salary_annuel > 0 THEN salary_annuel ELSE salaire END). Filtres : poste_actuel = TRUE, salaire > 0, non anonymises. Fourchette jauge : si >= 5 salaires -> [min x 0.9, max x 1.1] ; sinon [val x 0.7, val x 1.3] + mention 'indicatif'."),
        ("Tables / colonnes", "EXPERIENCE_PRO.salary_annuel, salaire, poste_actuel"),
        ("Exclusions / limites", "Salaires nuls exclus. Echantillon < 5 : fourchette elargie, mention affichee."),
        ("Exemple chiffre", "38000, 42000, 50000 -> moyenne 43333 ; n=3 < 5 -> fourchette [30333, 56333] indicatif"),
    ])

    section(doc, "2.10 Repartition par type de contrat")
    indicateur_table(doc, [
        ("Formule", "Par type : COUNT(*) des experiences en cours (poste_actuel = TRUE ou dates en cours)"),
        ("Tables / colonnes", "EXPERIENCE_PRO.type_contrat, poste_actuel, date_debut/fin"),
        ("Exclusions / limites", "On compte des experiences (pas des alumni). Valeurs vides -> 'Non renseigne'."),
        ("Exemple chiffre", "CDI 6, CDD 3, Stage 1, vide -> 'Non renseigne' 1"),
    ])

    # 3. Endpoints
    chapitre(doc, 3, "Endpoints API")
    endpoints = [
        ("GET /admin/indicateurs", "Indicateurs principaux (taux emploi 6 mois/global, taux_reponse, alumni_actifs, salaire_moyen/min/max, coherence, hypothese, source_de_verite)"),
        ("GET /admin/indicateurs/secteurs", "Repartition par secteur {secteur, count} + total_alumni"),
        ("GET /admin/indicateurs/types-contrat", "Repartition par type de contrat (vides -> Non renseigne)"),
        ("GET /admin/indicateurs/kpi-tag?tag=X", "Valeur d'un KPI specifique (valeur, unite, distribution, detail)"),
        ("GET /admin/indicateurs/kpi-tags", "Tous les tags KPI des questionnaires actifs, calcules automatiquement"),
        ("GET /admin/indicateurs/kpi-tags-actifs", "Tags DISTINCT des questionnaires actifs"),
        ("GET /admin/indicateurs/partenaires", "Indicateurs anonymises restreints aux alumni 'partage_donnees' actif"),
    ]
    styled_table(
        doc,
        headers=['Endpoint', 'Description'],
        rows=[(e, d) for e, d in endpoints],
        widths_cm=[5.0, 11.0],
    )

    # 4. Cas limites
    chapitre(doc, 4, "Cas limites & alertes")
    bullet(doc, "Promotion sans fenetre 6 mois ecoulee : taux null, statut en_attente, jamais comptee dans le global.")
    bullet(doc, "Ecart statut declaratif / poste reel mesure et expose (coherence_availability_poste_actuel).")
    bullet(doc, "Adequation sans reponse : etat vide avec instructions pour taguer une question KPI.")
    bullet(doc, "Echantillon salaire < 5 : fourchette elargie + mention 'echantillon limite'.")
    bullet(doc, "RGPD : comptes anonymises exclus ; salaire des anciens postes neutralise (salaire = 0).")

    out = os.path.join(OUTPUT_DIR, "methodologie_indicateurs_insertion.docx")
    doc.save(out)
    print("DOCX genere :", out)


if __name__ == "__main__":
    generate()
